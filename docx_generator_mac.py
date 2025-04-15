from docx import Document
from docx.shared import Inches, Pt
import os
from io import BytesIO
import matplotlib.pyplot as plt

def create_plot_image(data):
    fig, ax = plt.subplots(figsize=(4, 3))
    ax.plot(data)
    ax.set_title("Sample Chart")
    fig.tight_layout()

    image_stream = BytesIO()
    plt.savefig(image_stream, format='png', bbox_inches='tight')
    plt.close(fig)
    image_stream.seek(0)
    return image_stream

def insert_charts_into_template(template_path, output_docx_path, plots, comments):
    doc = Document(template_path)

    doc.add_heading("Generated Charts and Comments", level=1)

    table = doc.add_table(rows=3, cols=3)
    table.autofit = False

    plot_width = Inches(2.1)

    for i in range(9):
        row = i // 3
        col = i % 3
        cell = table.cell(row, col)

        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(plots[i], width=plot_width)

        p = cell.add_paragraph(comments[i])
        p.paragraph_format.space_after = 0
        p.paragraph_format.space_before = 0
        p.runs[0].font.size = Pt(8)

    doc.save(output_docx_path)

def main():
    os.makedirs("output", exist_ok=True)

    plots = [create_plot_image([j+i for j in range(5)]) for i in range(9)]
    comments = [f"This is the summary comment for chart {i+1}." for i in range(9)]

    template_path = "MC_template.docx"
    output_docx_path = "output/final_report_no_png.docx"

    insert_charts_into_template(template_path, output_docx_path, plots, comments)
    print(f"Report saved to: {output_docx_path}")

if __name__ == "__main__":
    main()



# from docx import Document
# from docx.shared import Inches, Pt
# import matplotlib.pyplot as plt
# import os
#
# def insert_charts_into_template(template_path, output_docx_path, plots, comments):
#     from docx.oxml.ns import qn
#     from docx.oxml import OxmlElement
#
#     doc = Document(template_path)
#     doc.add_heading("Generated Charts and Comments", level=1)
#
#     table = doc.add_table(rows=3, cols=3)
#     table.autofit = False
#
#     plot_width = Inches(2.1)
#
#     for i in range(9):
#         row = i // 3
#         col = i % 3
#         cell = table.cell(row, col)
#
#         paragraph = cell.paragraphs[0]
#         run = paragraph.add_run()
#         run.add_picture(plots[i], width=plot_width)
#
#         p = cell.add_paragraph(comments[i])
#         p.paragraph_format.space_after = 0
#         p.paragraph_format.space_before = 0
#
#         # Optional: shrink font size
#         run = p.runs[0]
#         run.font.size = Pt(8)
#
#     doc.save(output_docx_path)
#
# def generate_plot(data, filename):
#
#     plt.figure(figsize=(4, 3))
#     plt.plot(data)
#     plt.title("Sample Chart")
#     plt.tight_layout()
#     plt.savefig(filename)
#     plt.close()
#
# def main():
#     os.makedirs("plots", exist_ok=True)
#     os.makedirs("output", exist_ok=True)
#
#     plots = []
#     comments = [f"This is the summary comment for chart {i+1}." for i in range(9)]
#
#     for i in range(9):
#         path = f"plots/chart_{i+1}.png"
#         generate_plot([j+i for j in range(5)], path)
#         plots.append(path)
#
#     template_path = "MC_template.docx"
#     output_docx_path = "output/final_report.docx"
#
#     insert_charts_into_template(template_path, output_docx_path, plots, comments)
#     print(f"Report saved to: {output_docx_path}")
#
# if __name__ == "__main__":
#     main()